"""Document processing service using Azure Document Intelligence.

This module provides document parsing and processing capabilities using Azure Document Intelligence.
It handles various document formats including PDF, DOCX, Markdown, and plain text.

Example:
    ```python
    # Initialize service
    doc_service = DocumentService()
    
    # Parse a PDF file
    with open('document.pdf', 'rb') as f:
        content, metadata = doc_service.parse_file(f, 'application/pdf')
    ```
"""

import os
import logging
from functools import wraps
from typing import Dict, Any, List, Optional, BinaryIO, Tuple, Callable, TypeVar
from bs4 import BeautifulSoup
import docx
import markdown
from django.core.exceptions import ValidationError
from azure.core.exceptions import AzureError
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

logger = logging.getLogger(__name__)

def azure_retry(max_attempts: int = 3, min_wait: int = 4, max_wait: int = 10) -> Callable[[Callable[..., Any]], Callable[..., Any]]:
    """Decorator for retrying Azure operations with exponential backoff.
    
    This decorator provides resilient retry logic for Azure operations that may fail
    transiently. It uses exponential backoff to avoid overwhelming the service.
    
    Args:
        max_attempts (int): Maximum number of retry attempts. Defaults to 3.
        min_wait (int): Minimum wait time between retries in seconds. Defaults to 4.
        max_wait (int): Maximum wait time between retries in seconds. Defaults to 10.
        
    Returns:
        Callable: Decorated function with retry logic
    """
    def decorator(func: Callable[..., Any]) -> Callable[..., Any]:
        @retry(
            stop=stop_after_attempt(max_attempts),
            wait=wait_exponential(multiplier=1, min=min_wait, max=max_wait),
            retry=retry_if_exception_type(AzureError),
            reraise=True
        )
        @wraps(func)
        def wrapper(*args, **kwargs):
            return func(*args, **kwargs)
        return wrapper
    return decorator

class DocumentService:
    """Service for processing documents using Azure Document Intelligence.
    
    This service provides methods to parse and process various document formats using
    Azure Document Intelligence. It supports PDF, DOCX, Markdown, and plain text files.
    
    Attributes:
        doc_intelligence_client (DocumentIntelligenceClient): Azure Document Intelligence client
        
    Required Environment Variables:
        AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT: Document Intelligence endpoint
        AZURE_DOCUMENT_INTELLIGENCE_API_KEY: Document Intelligence API key
    """
    
    def __init__(self):
        """Initialize document service with Azure Document Intelligence client."""
        self.log_init("DocumentService")
        
        # Get configuration
        self.endpoint = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT')
        self.api_key = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_API_KEY')
        
        # Validate configuration
        if not self.endpoint or not self.api_key:
            error_msg = "Missing required environment variables for Document Intelligence"
            self.log_error(error_msg)
            raise ValueError(error_msg)
        
        # Initialize Document Intelligence client
        try:
            self.doc_intelligence_client = self.initialize_document_intelligence_client()
            self.log_success("Successfully initialized Document Intelligence client")
        except Exception as e:
            self.log_error("Failed to initialize Document Intelligence client", e)
            raise
    
    def log_init(self, service_name: str) -> None:
        """Log service initialization."""
        logger.info(f"Initializing {service_name}...")
    
    def log_success(self, message: str) -> None:
        """Log a success message."""
        logger.info(message)
    
    def log_error(self, message: str, error: Optional[Exception] = None) -> None:
        """Log an error message with optional exception."""
        if error:
            logger.error(f"{message}: {str(error)}")
        else:
            logger.error(message)
    
    def initialize_document_intelligence_client(self) -> DocumentIntelligenceClient:
        """Initialize Azure Document Intelligence client."""
        return DocumentIntelligenceClient(
            endpoint=self.endpoint,
            credential=AzureKeyCredential(self.api_key)
        )
    
    def parse_file(self, file_obj: BinaryIO, content_type: str) -> Tuple[str, Dict[str, Any]]:
        """Parse a document file and extract its content and metadata.
        
        Args:
            file_obj (BinaryIO): File-like object containing the document
            content_type (str): MIME type of the document. Supported types:
                - application/pdf
                - application/vnd.openxmlformats-officedocument.wordprocessingml.document
                - text/markdown
                - text/plain
                
        Returns:
            Tuple[str, Dict[str, Any]]: A tuple containing:
                - Extracted text content
                - Document metadata
                
        Raises:
            ValueError: If content type is not supported
            Exception: If parsing fails
        """
        try:
            if content_type == 'application/pdf':
                return self._parse_pdf(file_obj)
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._parse_docx(file_obj)
            elif content_type == 'text/markdown':
                return self._parse_markdown(file_obj)
            elif content_type == 'text/plain':
                return self._parse_text(file_obj)
            else:
                error_msg = f"Unsupported content type: {content_type}"
                self.log_error(error_msg)
                raise ValueError(error_msg)
        except Exception as e:
            self.log_error(f"Failed to parse file with content type {content_type}", e)
            raise
            
    @azure_retry(max_attempts=3, min_wait=2, max_wait=10)
    def _parse_pdf(self, file_obj: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF using Azure Document Intelligence.
        
        Args:
            file_obj (BinaryIO): PDF file object
            
        Returns:
            Tuple[str, Dict[str, Any]]: Extracted text and metadata
            
        Raises:
            Exception: If PDF parsing fails
        """
        try:
            poller = self.doc_intelligence_client.begin_analyze_document(
                "prebuilt-document",
                body=file_obj
            )
            result = poller.result()
            
            content = []
            for page in result.pages:
                if page.lines:
                    content.extend(line.content for line in page.lines)
                    
            metadata = {
                "page_count": len(result.pages),
                "language": result.language,
                "document_type": "pdf"
            }
            
            self.log_success(f"Successfully parsed PDF with {metadata['page_count']} pages")
            return "\n".join(content), metadata
            
        except Exception as e:
            self.log_error("Failed to parse PDF", e)
            raise
            
    @azure_retry(max_attempts=3, min_wait=2, max_wait=10)
    def _parse_docx(self, file_obj: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX file.
        
        Args:
            file_obj (BinaryIO): DOCX file object
            
        Returns:
            Tuple[str, Dict[str, Any]]: Extracted text and metadata
            
        Raises:
            Exception: If DOCX parsing fails
        """
        try:
            doc = docx.Document(file_obj)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
                    
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "document_type": "docx"
            }
            
            self.log_success(f"Successfully parsed DOCX with {metadata['paragraph_count']} paragraphs")
            return "\n".join(content), metadata
            
        except Exception as e:
            self.log_error("Failed to parse DOCX", e)
            raise
            
    @azure_retry(max_attempts=3, min_wait=2, max_wait=10)
    def _parse_markdown(self, file_obj: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """Parse Markdown file.
        
        Args:
            file_obj (BinaryIO): Markdown file object
            
        Returns:
            Tuple[str, Dict[str, Any]]: Raw markdown content and metadata
            
        Raises:
            Exception: If Markdown parsing fails
        """
        try:
            content = file_obj.read().decode('utf-8')
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            
            metadata = {
                "document_type": "markdown",
                "has_headers": bool(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']))
            }
            
            self.log_success("Successfully parsed Markdown")
            return content, metadata
            
        except Exception as e:
            self.log_error("Failed to parse Markdown", e)
            raise
            
    @azure_retry(max_attempts=3, min_wait=2, max_wait=10)
    def _parse_text(self, file_obj: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text file.
        
        Args:
            file_obj (BinaryIO): Text file object
            
        Returns:
            Tuple[str, Dict[str, Any]]: Text content and metadata
            
        Raises:
            Exception: If text parsing fails
        """
        try:
            content = file_obj.read().decode('utf-8')
            lines = content.splitlines()
            
            metadata = {
                "document_type": "text",
                "line_count": len(lines)
            }
            
            self.log_success(f"Successfully parsed text file with {metadata['line_count']} lines")
            return content, metadata
            
        except Exception as e:
            self.log_error("Failed to parse text file", e)
            raise
